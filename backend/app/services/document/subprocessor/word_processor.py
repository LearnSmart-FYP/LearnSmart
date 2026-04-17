import logging
import base64
import shutil
import subprocess
import tempfile
from pathlib import Path
from uuid import UUID

from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT
import fitz  # PyMuPDF
from lxml import etree

from app.services.document.subprocessor.table_processor import table_processor
from app.services.document.subprocessor.image_processor import image_processor
from app.services.document.extraction_tree import ExtractionTree

logger = logging.getLogger(__name__)

def convert_to_pdf_for_ocr(file_path: str) -> str:

    file_path = Path(file_path)

    with tempfile.TemporaryDirectory() as temp_dir:
        profile_dir = tempfile.mkdtemp(prefix="lo_profile_")
        try:
            result = subprocess.run(
                [
                    'libreoffice',
                    '--headless',
                    f'-env:UserInstallation=file://{profile_dir}',
                    '--convert-to', 'pdf',
                    '--outdir', temp_dir,
                    str(file_path)
                ],
                capture_output = True,
                text = True,
                timeout = 180)
        finally:
            shutil.rmtree(profile_dir, ignore_errors=True)

        if result.returncode != 0:
            raise RuntimeError(f"LibreOffice PDF conversion failed: {result.stderr}")

        converted_path = Path(temp_dir) / f"{file_path.stem}.pdf"

        if not converted_path.exists():
            raise RuntimeError(f"Converted PDF not found: {converted_path}")

        final_path = Path(tempfile.gettempdir()) / f"{file_path.stem}_ocr_{id(file_path)}.pdf"
        final_path.write_bytes(converted_path.read_bytes())

        return str(final_path)

def convert_doc_to_docx(doc_path: str) -> str:

    doc_path = Path(doc_path)

    with tempfile.TemporaryDirectory() as temp_dir:

        profile_dir = tempfile.mkdtemp(prefix="lo_profile_")
        try:
            result = subprocess.run(
                [
                    'libreoffice',
                    '--headless',
                    f'-env:UserInstallation=file://{profile_dir}',
                    '--convert-to', 'docx',
                    '--outdir', temp_dir,
                    str(doc_path)
                ],
                capture_output = True,
                text = True,
                timeout = 120)
        finally:
            shutil.rmtree(profile_dir, ignore_errors = True)

        if result.returncode != 0:
            raise RuntimeError(f"LibreOffice conversion failed: {result.stderr}")

        # Find the converted file
        converted_path = Path(temp_dir) / f"{doc_path.stem}.docx"

        if not converted_path.exists():
            raise RuntimeError(f"Converted file not found: {converted_path}")

        final_path = doc_path.with_suffix('.docx')
        final_path.write_bytes(converted_path.read_bytes())

        logger.info(f"Converted {doc_path.name} to {final_path.name}")
        return str(final_path)

class WordProcessor:

    def __init__(self):

        # Use global helper processors
        self.image_processor = image_processor
        self.table_processor = table_processor

    def _get_full_document_ocr(self, file_path: str) -> str:

        ocr_texts = []
        pdf_path = None

        try:

            logger.info(f"Converting to PDF for full-page OCR")
            pdf_path = convert_to_pdf_for_ocr(file_path)
            doc = fitz.open(pdf_path)

            for i, page in enumerate(doc):
                page_number = i + 1
                try:
                    mat = fitz.Matrix(2, 2)
                    pix = page.get_pixmap(matrix = mat)
                    page_image_bytes = pix.tobytes("png")

                    ocr_result = self.image_processor.process(page_image_bytes)
                    if ocr_result.success and ocr_result.output_text:
                        ocr_texts.append(f"[Page {page_number}]\n{ocr_result.output_text}")
                except Exception as e:
                    logger.debug(f"Full-page OCR failed for page {page_number}: {e}")

            doc.close()
            logger.info(f"Full-page OCR completed for {len(ocr_texts)} pages")

        except Exception as e:
            logger.warning(f"PDF conversion for OCR failed: {e}")

        finally:
            if pdf_path and Path(pdf_path).exists():
                try:
                    Path(pdf_path).unlink()
                except Exception:
                    pass

        return "\n\n".join(ocr_texts)

    def _extract_comments(self, doc: Document) -> list[dict]:

        comments = []
        nsmap = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}

        try:
            if hasattr(doc.part, 'rels'):
                for rel in doc.part.rels.values():
                    if 'comments' in rel.reltype.lower() and 'extended' not in rel.reltype.lower():
                        try:

                            comments_part = rel.target_part
                            root = etree.fromstring(comments_part.blob)

                            for comment in root.findall('.//w:comment', nsmap):
                                author = comment.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}author', '')
                                date = comment.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}date', '')
                                text_parts = []
                                for t in comment.iter('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t'):
                                    if t.text:
                                        text_parts.append(t.text)
                                text = ''.join(text_parts).strip()
                                if text:
                                    comments.append({
                                        'author': author,
                                        'date': date,
                                        'text': text})
                                    
                        except Exception as e:
                            logger.debug(f"Comments part extraction failed: {e}")

        except Exception as e:
            logger.debug(f"Comments extraction failed: {e}")

        return comments

    def process_sync(
        self,
        file_path: str,
        tree: ExtractionTree | None = None,
        parent_node_id: UUID | None = None) -> ExtractionTree:

        total_tables = 0
        total_images = 0
        total_links = 0
        total_embedded = 0
        converted_file = None

        try:

            # Convert .doc to .docx if necessary

            file_path_obj = Path(file_path)
            if file_path_obj.suffix.lower() == '.doc':
                logger.info(f"Converting legacy .doc file: {file_path}")
                converted_file = convert_doc_to_docx(file_path)
                file_path = converted_file

            doc = Document(file_path)
            doc_base_name = Path(file_path).stem
            full_ocr_text = self._get_full_document_ocr(file_path)

            # Create or use existing tree

            if tree is None:
                tree = ExtractionTree()
                root_id = tree.add_root(
                    source_path = file_path,
                    document_type = "word",
                    metadata = {})
            else:
                root_id = tree.add_child(
                    parent_id = parent_node_id,
                    source_path = file_path,
                    document_type = "word",
                    extraction_type = "embedded_document",
                    metadata = {})

            # Extract paragraphs
            paragraphs = []
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    paragraphs.append(text)

            # Add text content node
            if paragraphs:
                tree.add_child(
                    parent_id = root_id,
                    source_path = f"{file_path}:text",
                    document_type = "text",
                    extraction_type = "text_content",
                    metadata = {"output_text": "\n\n".join(paragraphs)})

            # Extract tables
            for table_idx, table in enumerate(doc.tables):
                try:
                    table_data = []

                    for row in table.rows:
                        row_data = [cell.text.strip() for cell in row.cells]
                        table_data.append(row_data)

                    if not table_data:
                        continue

                    table_result = self.table_processor.process(
                        table_data = table_data)

                    if table_result.success:
                        tree.add_child(
                            parent_id = root_id,
                            source_path = f"{file_path}:table{table_idx}",
                            document_type = "table",
                            extraction_type = "table",
                            metadata = {"output_text": table_result.output_text})
                        total_tables += 1

                except Exception as e:
                    logger.debug(f"Table {table_idx} extraction failed: {e}")

            # Extract images
            for rel_id, rel in doc.part.rels.items():
                if "image" in rel.reltype:
                    try:

                        image_part = rel.target_part
                        image_bytes = image_part.blob
                        image_ext = rel.target_ref.split('.')[-1] if '.' in rel.target_ref else 'png'

                        # Run OCR on image bytes
                        output_text = ""
                        try:
                            ocr_result = self.image_processor.process(image_bytes)
                            output_text = ocr_result.output_text if ocr_result.success else ""
                        except Exception:
                            output_text = ""

                        tree.add_child(
                            parent_id = root_id,
                            source_path = f"{file_path}:image{total_images}",
                            document_type = "image",
                            extraction_type = "image",
                            metadata = {
                                "data": base64.b64encode(image_bytes).decode('utf-8'),
                                "format": image_ext,
                                "output_text": output_text,
                                "suggested_filename": f"{doc_base_name}_img{total_images}.{image_ext}"})
                        total_images += 1

                    except Exception as e:
                        logger.debug(f"Image extraction failed: {e}")

            # Extract hyperlinks
            for rel_id, rel in doc.part.rels.items():
                if rel.reltype == RT.HYPERLINK:
                    try:

                        uri = rel.target_ref
                        if not uri:
                            continue

                        link_type = "external" if uri.startswith(('http://', 'https://')) else "internal"

                        tree.add_child(
                            parent_id = root_id,
                            source_path = f"{file_path}:link{total_links}",
                            document_type = "link",
                            extraction_type = "link",
                            metadata = {
                                "uri": uri,
                                "link_type": link_type})
                        total_links += 1

                    except Exception as e:
                        logger.debug(f"Link extraction failed: {e}")

            # Extract comments (kept because comments may be hidden in print/PDF view)
            comments = self._extract_comments(doc)
            if comments:
                tree.add_child(
                    parent_id = root_id,
                    source_path = f"{file_path}:comments",
                    document_type = "text",
                    extraction_type = "comments",
                    metadata = {
                        "comments": comments,
                        "output_text": "\n\n".join(
                            f"[{c['author']}] {c['text']}" for c in comments),
                        "count": len(comments)})

            if full_ocr_text:
                tree.add_child(
                    parent_id = root_id,
                    source_path = f"{file_path}:full_ocr",
                    document_type = "text",
                    extraction_type = "full_ocr",
                    metadata = {
                        "output_text": f"[Full Document OCR]:\n{full_ocr_text}"})

            # Extract embedded files (collected for later processing)
            for rel_id, rel in doc.part.rels.items():
                rel_type_lower = rel.reltype.lower() if rel.reltype else ""
                if "oleobject" in rel_type_lower or "package" in rel_type_lower:
                    try:

                        emb_part = rel.target_part
                        emb_bytes = emb_part.blob
                        emb_filename = f"{doc_base_name}_embedded{total_embedded}.bin"

                        tree.add_child(
                            parent_id = root_id,
                            source_path = f"{file_path}:embedded:{emb_filename}",
                            document_type = "embedded",
                            extraction_type = "embedded_document",
                            metadata = {
                                "data": base64.b64encode(emb_bytes).decode('utf-8'),
                                "file_size": len(emb_bytes),
                                "suggested_filename": emb_filename})

                        total_embedded += 1
                        logger.debug(f"Extracted embedded file: {emb_filename}")

                    except Exception as e:
                        logger.debug(f"Embedded file extraction failed: {e}")

            logger.info(f"Word processed: {len(paragraphs)} paragraphs, "
                       f"Tables: {total_tables}, Images: {total_images}, "
                       f"Links: {total_links}, Embedded: {total_embedded}, "
                       f"Comments: {len(comments)}, "
                       f"Full OCR: {'Yes' if full_ocr_text else 'No'}")

            return tree

        except Exception as e:
            logger.error(f"Word processing failed: {e}")
            raise

        finally:

            if converted_file and Path(converted_file).exists():
                try:
                    Path(converted_file).unlink()
                except Exception:
                    pass

# Global instance
word_processor = WordProcessor()
